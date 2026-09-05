#!/usr/bin/env python3
"""Audit the JAIS submission-facing files for final editorial hygiene."""

from __future__ import annotations

import json
import zipfile
from collections import OrderedDict
from pathlib import Path

from docx import Document

HERE = Path(__file__).resolve().parent
GENERATED = HERE / "upload-packet" / "generated"
DOCX = GENERATED / "JAIS_MANUSCRIPT.docx"
OUTPUT_JSON = GENERATED / "JAIS_SUBMISSION_STYLE_AUDIT.json"
OUTPUT_MD = GENERATED / "JAIS_SUBMISSION_STYLE_AUDIT.md"

TEXT_TARGETS = [
    HERE / "cover-letter.md",
    HERE / "ai-disclosure.md",
    HERE / "title-page.md",
    HERE / "jais-abstract.md",
]

EM_DASH = "\u2014"
OLD_TARGET_RESIDUE = (
    "primary target remains Computers & Security",
    "for consideration by Computers & Security",
    "submitted to Computers & Security",
    "Computers & Security submission package",
)


def iter_paragraphs(container):
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def docx_text(path: Path) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in iter_paragraphs(doc))


def inspect_ooxml(path: Path) -> tuple[bool, bool]:
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        has_comments = "word/comments.xml" in names
        document_xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
        has_tracked = "<w:ins " in document_xml or "<w:del " in document_xml
    return has_comments, has_tracked


def main() -> int:
    manuscript = docx_text(DOCX)
    has_comments, has_tracked = inspect_ooxml(DOCX)
    source_text = {p.name: p.read_text(encoding="utf-8") for p in TEXT_TARGETS}

    checks = OrderedDict()
    checks["manuscript_no_em_dash"] = EM_DASH not in manuscript
    checks["submission_text_files_no_em_dash"] = all(EM_DASH not in text for text in source_text.values())
    checks["cover_letter_uses_exact_full_paper_label"] = "Full Paper" in source_text["cover-letter.md"]
    checks["cover_letter_has_no_old_article_label"] = "Regular/Full Article" not in source_text["cover-letter.md"]
    checks["title_page_uses_exact_full_paper_label"] = "Full Paper" in source_text["title-page.md"]
    checks["title_page_has_no_old_article_label"] = "Regular/Full Article" not in source_text["title-page.md"]
    checks["ai_disclosure_names_tool_transparently"] = "OpenAI ChatGPT" in source_text["ai-disclosure.md"]
    checks["manuscript_ai_disclosure_present"] = "OpenAI ChatGPT" in manuscript
    checks["no_tracked_changes"] = not has_tracked
    checks["no_word_comments"] = not has_comments
    checks["no_old_target_submission_language"] = not any(phrase in manuscript for phrase in OLD_TARGET_RESIDUE)

    result = {
        "checks": checks,
        "style_gate_pass": all(checks.values()),
        "cover_letter_words": len(source_text["cover-letter.md"].split()),
        "ai_disclosure_words": len(source_text["ai-disclosure.md"].split()),
        "manuscript_em_dash_count": manuscript.count(EM_DASH),
        "tracked_changes_present": has_tracked,
        "comments_present": has_comments,
    }

    OUTPUT_JSON.write_text(json.dumps(result, indent=2) + "\n", encoding="utf-8")

    lines = [
        "# JAIS Submission-Facing Style Audit",
        "",
        f"- Cover-letter words: {result['cover_letter_words']}",
        f"- AI-disclosure words: {result['ai_disclosure_words']}",
        f"- Manuscript em-dash count: {result['manuscript_em_dash_count']}",
        f"- Tracked changes present: {result['tracked_changes_present']}",
        f"- Word comments present: {result['comments_present']}",
        "",
        "## Checks",
        "",
    ]
    for name, passed in checks.items():
        lines.append(f"- [{'x' if passed else ' '}] {name}")
    lines.extend([
        "",
        f"**Submission-facing style gate:** {'PASS' if result['style_gate_pass'] else 'FAIL'}",
        "",
        "This audit checks editorial hygiene only. It does not replace the frozen-science, reference, provenance, or ScholarOne attestation gates.",
    ])
    OUTPUT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")

    for name, passed in checks.items():
        print(f"{name}={passed}")
    print(f"submission_style_gate_pass={result['style_gate_pass']}")
    print(f"manuscript_em_dash_count={result['manuscript_em_dash_count']}")
    print(f"tracked_changes_present={has_tracked}")
    print(f"comments_present={has_comments}")
    return 0 if result["style_gate_pass"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
