#!/usr/bin/env python3
"""Render the generated JAIS Markdown manuscript to an AIAA-style DOCX.

Requires python-docx. Run build_jais_final_export.py first or use --build.
"""

from __future__ import annotations

import argparse
import re
import subprocess
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError as exc:
    raise SystemExit("python-docx is required to generate the JAIS Word manuscript") from exc

HERE = Path(__file__).resolve().parent
SOURCE = HERE / "upload-packet" / "generated" / "JAIS_MANUSCRIPT.md"
OUTPUT = HERE / "upload-packet" / "generated" / "JAIS_MANUSCRIPT.docx"
BLACK = RGBColor(0, 0, 0)


def clean_inline(text: str) -> str:
    """Remove Markdown-only inline delimiters that should not appear in Word."""
    return text.replace("`", "")


def style_run(run, size: float = 10, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    if bold is not None:
        run.bold = bold


def set_cell_text(cell, text: str, header: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(clean_inline(text).replace("\\|", "|"))
    style_run(r, 8, header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    style_run(run, 10)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_markdown_paragraph(
    doc: Document,
    text: str,
    style=None,
    align=None,
    line_spacing: float = 2.0,
    size: float = 10,
):
    text = clean_inline(text)
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(0)

    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > cursor:
            r = p.add_run(text[cursor:match.start()])
            style_run(r, size)
        r = p.add_run(match.group(1))
        style_run(r, size, True)
        cursor = match.end()
    if cursor < len(text):
        r = p.add_run(text[cursor:])
        style_run(r, size)
    return p


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def column_widths(rows: list[list[str]], total_inches: float = 6.5) -> list[float]:
    cols = max(len(row) for row in rows)
    max_lengths = []
    for col in range(cols):
        values = [len(row[col]) if col < len(row) else 0 for row in rows]
        max_lengths.append(max(8, min(max(values), 80)))
    weights = [max(0.75, length ** 0.55) for length in max_lengths]
    scale = total_inches / sum(weights)
    return [weight * scale for weight in weights]


def build_docx(source: Path, output: Path) -> None:
    text = source.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 2.0
    styles["Heading 1"].font.size = Pt(11)
    styles["Heading 2"].font.size = Pt(10)
    styles["Heading 3"].font.size = Pt(10)

    add_page_number(section.footer.paragraphs[0])

    i = 0
    first_heading = True
    frontmatter = True
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue

        if line.startswith("|"):
            block = []
            while i < len(lines) and lines[i].startswith("|"):
                block.append(lines[i])
                i += 1
            parsed = []
            for row in block:
                cells = [c.strip() for c in row.strip().strip("|").split("|")]
                parsed.append(cells)
            if len(parsed) >= 2 and all(
                re.fullmatch(r"-+", c.replace(":", "")) for c in parsed[1]
            ):
                parsed.pop(1)
            if parsed:
                cols = max(len(row) for row in parsed)
                table = doc.add_table(rows=len(parsed), cols=cols)
                table.style = "Table Grid"
                table.autofit = False
                widths = column_widths(parsed)
                for rr, row in enumerate(parsed):
                    set_row_cant_split(table.rows[rr])
                    if rr == 0:
                        set_repeat_table_header(table.rows[rr])
                    for cc in range(cols):
                        cell = table.cell(rr, cc)
                        set_cell_text(cell, row[cc] if cc < len(row) else "", rr == 0)
                        cell.width = Inches(widths[cc])
                doc.add_paragraph()
            continue

        if line.startswith("# "):
            title = line[2:].strip()
            if first_heading:
                # Use a plain paragraph instead of Word's themed Title style so
                # no accent-color border/theme formatting leaks into submission.
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(8)
                r = p.add_run(title)
                style_run(r, 12, True)
                first_heading = False
            else:
                add_markdown_paragraph(doc, title, "Heading 1")
            i += 1
            continue

        if line.startswith("## "):
            heading = line[3:].strip()
            if heading == "Abstract":
                frontmatter = False
            add_markdown_paragraph(doc, heading, "Heading 2")
            i += 1
            continue

        if line.startswith("### "):
            add_markdown_paragraph(doc, line[4:].strip(), "Heading 3")
            i += 1
            continue

        if frontmatter:
            add_markdown_paragraph(
                doc,
                line,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                line_spacing=1.0,
                size=10,
            )
            i += 1
            continue

        if line.startswith("- "):
            p = add_markdown_paragraph(doc, line[2:].strip(), "List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue

        if re.match(r"^\d+\.\s", line):
            p = add_markdown_paragraph(
                doc,
                re.sub(r"^\d+\.\s+", "", line),
                "List Number",
            )
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue

        para = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
            r"^(#{1,3}\s|\|\s|-\s|\d+\.\s)", lines[i]
        ):
            para.append(lines[i].strip())
            i += 1
        add_markdown_paragraph(doc, " ".join(para))

    doc.core_properties.title = (
        "Satellite Cyber Response and Trusted Recovery Under Contact and Adversarial Evidence Constraints"
    )
    doc.core_properties.subject = "AIAA Journal of Aerospace Information Systems submission manuscript"
    doc.core_properties.author = "Aman Kumar Singh"
    doc.core_properties.keywords = (
        "satellite cybersecurity, mission-aware cybersecurity, cyber resilience, trusted recovery"
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--build",
        action="store_true",
        help="run build_jais_final_export.py first",
    )
    args = parser.parse_args()

    if args.build or not SOURCE.exists():
        subprocess.run(
            [sys.executable, str(HERE / "build_jais_final_export.py")],
            check=True,
        )

    build_docx(SOURCE, OUTPUT)
    print(f"docx={OUTPUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
