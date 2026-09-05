#!/usr/bin/env python3
"""Apply submission-facing prose normalization to the generated JAIS manuscript.

This is a venue-facing editorial layer only. It does not alter frozen numerical
results, statistical populations, treatment identities, or claim boundaries.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

HERE = Path(__file__).resolve().parent
GENERATED = HERE / "upload-packet" / "generated"
MARKDOWN = GENERATED / "JAIS_MANUSCRIPT.md"
DOCX = GENERATED / "JAIS_MANUSCRIPT.docx"

EM_DASH = "\u2014"

TARGETED_REPLACEMENTS = {
    "The narrower anticipated mechanism\u2014nominal behavioral restoration without sufficient verification\u2014was not observed in A13.":
        "The narrower anticipated mechanism, nominal behavioral restoration without sufficient verification, was not observed in A13.",
    "During preparation of this journal work, the author used OpenAI ChatGPT to assist with manuscript organization, source checking, editorial refinement, consistency review, reproducibility documentation, repository and audit workflow support, and preparation of journal-submission materials. The author reviewed and edited all resulting content, checked scientific quantities and source claims against the frozen research record and cited sources, and takes full responsibility for the manuscript. For Study 1, this assistance occurred after the experimental campaign and historical statistical findings were frozen and after the evidence package had been archived. For Study 2, the assistance occurred after the campaign evidence and prospective analysis implementation were frozen. It did not generate or replace observations, alter seeds or exclusions, change either frozen statistical population, modify the frozen Study-2 analyzer, or provide input to the evaluated deterministic response policies.":
        "During preparation of this article, I used OpenAI ChatGPT for editorial and research-support tasks, including manuscript organization, source checking, language refinement, consistency review, reproducibility documentation, repository and audit workflow support, and preparation of submission materials. I reviewed and revised the resulting text, checked scientific quantities and source claims against the frozen research record and cited sources, and take full responsibility for the manuscript. For Study 1, this assistance was used only after the experimental campaign and historical statistical findings had been frozen and after the evidence package had been archived. For Study 2, it was used only after the campaign evidence and prospective analysis implementation had been frozen. It did not generate or replace observations, alter seeds or exclusions, change either frozen statistical population, modify the frozen Study 2 analyzer, or provide input to the evaluated deterministic response policies.",
}


def normalize_text(text: str) -> str:
    for old, new in TARGETED_REPLACEMENTS.items():
        text = text.replace(old, new)
    text = text.replace(" \u2014 ", ": ")
    text = text.replace(EM_DASH, "-")
    return text


def iter_paragraphs(container):
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def normalize_docx(path: Path) -> int:
    doc = Document(path)
    changed = 0
    for paragraph in iter_paragraphs(doc):
        replacement = normalize_text(paragraph.text)
        if replacement == paragraph.text:
            continue
        if paragraph.runs:
            paragraph.runs[0].text = replacement
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(replacement)
        changed += 1

    doc.core_properties.author = "Aman Kumar Singh"
    doc.core_properties.last_modified_by = "Aman Kumar Singh"
    doc.save(path)
    return changed


def extract_docx_text(path: Path) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in iter_paragraphs(doc))


def main() -> int:
    if not MARKDOWN.exists() or not DOCX.exists():
        raise SystemExit("Run the JAIS build and Word renderer before finalization")

    md_before = MARKDOWN.read_text(encoding="utf-8")
    md_after = normalize_text(md_before)
    MARKDOWN.write_text(md_after, encoding="utf-8")

    changed_paragraphs = normalize_docx(DOCX)
    docx_text = extract_docx_text(DOCX)

    if EM_DASH in md_after or EM_DASH in docx_text:
        raise SystemExit("Submission finalization failed: em dash remains in manuscript")

    print(f"markdown_em_dashes_removed={md_before.count(EM_DASH)}")
    print(f"docx_paragraphs_normalized={changed_paragraphs}")
    print("submission_punctuation_gate=PASS")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
